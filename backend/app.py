
import io
import re
import os
import json
import tempfile
import uvicorn
import textstat
from collections import Counter
from typing import Optional, List, Dict

# Core Libraries
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document

# PDF Generation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Conditional imports with fallbacks
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    print("✅ spaCy loaded successfully")
except:
    print("⚠️ spaCy not available - using fallback keyword extraction")
    nlp = None

# --- 1. UPDATED IMPORTS ---
try:
    from pypdf import PdfReader
    print("✅ pypdf (v6+) loaded successfully")
except ImportError:
    PdfReader = None
    print("⚠️ pypdf not found - run pip install pypdf")

# --- 2. FIXED EXTRACTION FUNCTION ---
def extract_pdf_text(file_obj):
    if PdfReader is None:
        return "PDF Error: pypdf library not installed on server."
    
    try:
        # file_obj is the io.BytesIO(file_content) passed from BotAudit
        reader = PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        
        if not text.strip():
            return "CRITICAL_ERROR_IMAGE_ONLY"
            
        return text
    except Exception as e:
        print(f"Extraction failed: {str(e)}")
        return f"PDF Error: {str(e)}"

# --- BOT AUDIT CLASS (SINGLE DEFINITION) ---
class BotAudit:
    """2026 Enterprise-Grade Resume Analysis Engine"""
    
    def __init__(self, nlp_model=None):
        self.nlp = nlp_model
        self.fluff_words = [
            "hardworking", "team player", "results-driven", "passionate", 
            "detail-oriented", "self-motivated", "synergy", "think outside the box",
            "go-getter", "people person", "perfectionist", "innovative"
        ]
        self.leadership_verbs = [
            "spearheaded", "negotiated", "orchestrated", "architected", 
            "transformed", "steered", "pioneered", "revolutionized", 
            "optimized", "streamlined", "accelerated", "dominated",
            "launched", "established", "founded", "directed", "managed"
        ]
        self.support_verbs = [
            "helped", "assisted", "supported", "participated", "contributed",
            "worked on", "was responsible for", "handled", "did", "performed"
        ]

    def parse_and_audit_structure(self, file_content: bytes, filename: str):
        """Multi-Stage Document Parser with Error Handling"""
        results = {
            "contact_in_header": False,
            "table_detected": False,
            "text": "",
            "parsing_safe": True,
            "file_type": filename.split('.')[-1].lower() if '.' in filename else "unknown"
        }
        
        try:
            if filename.lower().endswith('.docx'):
                doc = Document(io.BytesIO(file_content))
                
                # Header/Footer Contact Check
                contact_pattern = r'[\d\-\(\)\+]{10,}|@[\w\.-]+\.[a-zA-Z]{2,}|linkedin\.com'
                
                for section in doc.sections:
                    # Check headers
                    if section.header and section.header.paragraphs:
                        for para in section.header.paragraphs:
                            if re.search(contact_pattern, para.text, re.IGNORECASE):
                                results["contact_in_header"] = True

                # Table Detection
                if len(doc.tables) > 0:
                    results["table_detected"] = True

                # Extract Text
                full_text = ""
                for para in doc.paragraphs:
                    full_text += para.text + "\n"
                
                results["text"] = full_text
                
            elif filename.lower().endswith('.pdf'):
                try:
                    results["text"] = extract_pdf_text(io.BytesIO(file_content))
                    
                    # OCR Check
                    if len(results["text"].strip()) < 100:
                        results["text"] = "CRITICAL_ERROR_IMAGE_ONLY"
                        results["parsing_safe"] = False
                        
                except Exception as pdf_error:
                    print(f"PDF parsing error: {pdf_error}")
                    results["text"] = "PDF_PROCESSING_FAILED"
                    results["parsing_safe"] = False
                    
            else:
                # Plain text fallback
                try:
                    results["text"] = file_content.decode('utf-8', errors='ignore')
                except:
                    results["text"] = str(file_content, errors='ignore')
                
        except Exception as e:
            print(f"File parsing error: {e}")
            results["text"] = f"PARSING_FAILED: {str(e)}"
            results["parsing_safe"] = False
            
        return results

    def check_bot_safety(self, text: str, structure_audit: Dict = None):
        """Critical Bot-Beater Analysis"""
        warnings = []
        safety_score = 100
        
        # Critical Errors
        if text.startswith("CRITICAL_ERROR_IMAGE_ONLY"):
            return {
                "safety_score": 0,
                "is_bot_readable": False,
                "critical_warnings": ["🚨 CRITICAL: PDF is an image. ATS sees 0 words. Convert to text-based PDF."],
                "layout_risk": "CRITICAL"
            }
            
        if text.startswith("PDF_PROCESSING_FAILED"):
            return {
                "safety_score": 20,
                "is_bot_readable": False,
                "critical_warnings": ["🚨 PDF processing failed. Try converting to .docx format."],
                "layout_risk": "HIGH"
            }
            
        if text.startswith("PARSING_FAILED"):
            return {
                "safety_score": 0,
                "is_bot_readable": False,
                "critical_warnings": ["🚨 File parsing failed. Check file format and try again."],
                "layout_risk": "CRITICAL"
            }

        # Layout Issues
        if "\t" in text or re.search(r' {5,}', text):
            warnings.append("🚨 Multi-column layout detected. ATS may scramble text order.")
            safety_score -= 25
            
        if structure_audit and structure_audit.get("table_detected"):
            warnings.append("🚨 Tables detected. Convert to simple bullet points.")
            safety_score -= 30
            
        if structure_audit and structure_audit.get("contact_in_header"):
            warnings.append("🚨 Contact info in header/footer is INVISIBLE to ATS bots.")
            safety_score -= 25

        # Section Headers Check
        standard_sections = ["experience", "education", "skills", "summary"]
        found_sections = [s for s in standard_sections if s in text.lower()]
        
        if len(found_sections) < 3:
            warnings.append("🚨 Missing standard section headers. Use: Experience, Education, Skills")
            safety_score -= 15
            
        return {
            "safety_score": max(safety_score, 0),
            "is_bot_readable": safety_score > 70,
            "critical_warnings": warnings,
            "layout_risk": "HIGH" if safety_score < 50 else "MEDIUM" if safety_score < 80 else "LOW",
            "found_sections": found_sections
        }

    def analyze_advanced_metrics(self, text: str):
        """Advanced Resume Metrics Analysis"""
        if text.startswith(("CRITICAL_ERROR", "PDF_PROCESSING_FAILED", "PARSING_FAILED")):
            return {
                "readability": {"grade_level": 0, "assessment": "Cannot analyze - file processing failed"},
                "verb_strength": 0,
                "metric_count": 0,
                "stuffed_keywords": [],
                "fluff_words": [],
                "impact_score": 0,
                "analysis_possible": False
            }
        
        lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 10]
        text_lower = text.lower()
        
        # 1. Readability Analysis
        try:
            grade = textstat.flesch_kincaid_grade(text)
            assessment = "Easy to Read" if grade < 10 else "Professional Level" if grade < 14 else "Too Complex"
        except:
            grade = 12
            assessment = "Professional Level"
            
        # 2. Action Verb Analysis
        power_count = sum(1 for v in self.leadership_verbs if v in text_lower)
        weak_count = sum(1 for v in self.support_verbs if v in text_lower)
        
        if power_count + weak_count > 0:
            verb_strength = (power_count / (power_count + weak_count)) * 100
        else:
            verb_strength = 0
        
        # 3. Quantifiable Metrics
        metric_patterns = [r'\d+%', r'\$[\d,]+', r'\d+\+', r'\d+x', r'\d+k\+?', r'\d+m\+?']
        metric_count = sum(len(re.findall(pattern, text, re.IGNORECASE)) for pattern in metric_patterns)
        
        # 4. Keyword Stuffing Detection
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text_lower)
        if words:
            word_freq = Counter(words)
            total_words = len(words)
            stuffed = [(w, c, round(c/total_words*100, 1)) for w, c in word_freq.items() 
                      if c/total_words > 0.05 and c > 3]
        else:
            stuffed = []

        # 5. Buzzword Detection
        fluff_found = [word for word in self.fluff_words if word in text_lower]

        # 6. Digital Presence
        linkedin_found = bool(re.search(r'linkedin\.com/in/[\w\-]+', text_lower))
        email_found = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text))

        return {
            "readability": {"grade_level": grade, "assessment": assessment},
            "verb_strength": round(verb_strength, 1),
            "metric_count": metric_count,
            "stuffed_keywords": stuffed,
            "fluff_words": fluff_found,
            "impact_score": min(100, metric_count * 20),
            "digital_presence": {"linkedin_found": linkedin_found, "email_found": email_found},
            "analysis_possible": True
        }

    def extract_keywords(self, text: str, max_keywords: int = 15):
        """Keyword Extraction with Fallback"""
        if self.nlp:
            try:
                doc = self.nlp(text.lower())
                keywords = []
                for token in doc:
                    if (token.pos_ in ['NOUN', 'PROPN', 'ADJ'] and 
                        not token.is_stop and 
                        not token.is_punct and 
                        len(token.text) > 3):
                        keywords.append(token.text)
                
                keyword_counts = Counter(keywords)
                return [kw for kw, count in keyword_counts.most_common(max_keywords)]
            except:
                pass
        
        # Fallback method
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        stop_words = {
            'that', 'with', 'have', 'this', 'will', 'from', 'they', 'been', 'were',
            'when', 'what', 'where', 'there', 'these', 'those', 'then', 'than'
        }
        filtered = [w for w in words if w not in stop_words]
        word_counts = Counter(filtered)
        return [word for word, count in word_counts.most_common(max_keywords)]

# --- FASTAPI APP (SINGLE DEFINITION) ---
app = FastAPI(
    title="Trendsetter AI Resume Helper 2026",
    description="Enterprise-Grade Resume Optimization Engine - Backend for https://trendsetter-resume-helper.onrender.com",
    version="2.0.0"
)

# ...existing code...


app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://trendsetter-resume-helper.onrender.com",  # Production backend
        "http://localhost:3000",                          # Local development
        "http://127.0.0.1:3000",
        "http://localhost:5500",                          # Live Server
        "http://127.0.0.1:5500",
        "https://*.vercel.app",                           # Vercel deployments wildcard
        "https://trendsetter-resume-helper.vercel.app",   # Specific Vercel URL
        "*"                                               # Development fallback
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 
# Initialize Bot Audit AFTER class definition
bot_audit = BotAudit(nlp)

# --- API ENDPOINTS ---

@app.get("/")
async def root():
    return {
        "message": "Trendsetter AI Resume Helper 2026",
        "status": "active",
        "version": "2.0.0",
        "features": {
            "spacy_nlp": nlp is not None,
            "pdf_processing": "extract_pdf_text" in globals(),
            "docx_processing": True,
            "advanced_analytics": True
        }
    }

@app.post("/api/scan-resume")
async def scan_resume(file: UploadFile = File(...)):
    try:
        content = await file.read()
        
        # Parse and audit file structure
        audit_results = bot_audit.parse_and_audit_structure(content, file.filename)
        
        # Analyze bot safety
        bot_safety = bot_audit.check_bot_safety(audit_results["text"], audit_results)
        
        # Advanced metrics analysis
        metrics = bot_audit.analyze_advanced_metrics(audit_results["text"])
        
        return {
            "success": True,
            "filename": file.filename,
            "file_type": audit_results["file_type"],
            "text_length": len(audit_results["text"]),
            "parsing_successful": audit_results["parsing_safe"],
            "bot_safety": bot_safety,
            "metrics": metrics,
            "text_preview": audit_results["text"][:300] + "..." if len(audit_results["text"]) > 300 else audit_results["text"]
        }
        
    except Exception as e:
        print(f"Scan error: {e}")
        raise HTTPException(status_code=500, detail=f"File processing error: {str(e)}")

# ...existing code...

# ...existing code...

# ...existing code...

@app.post("/api/match-job")
async def match_job(
    resume_text: str = Form(...),
    job_description: str = Form(...),
    job_title: str = Form(""),
    resume_file: UploadFile = File(None)
):
    try:
        print(f"🔍 Processing job match for: {job_title or 'Unknown Position'}")
        
        # Process uploaded file if provided
        structure_audit = None
        if resume_file and resume_file.filename:
            print("📄 Processing uploaded resume file...")
            file_content = await resume_file.read()
            parsed_result = bot_audit.parse_and_audit_structure(file_content, resume_file.filename)
            if parsed_result["parsing_safe"]:
                resume_text = parsed_result["text"]
                structure_audit = parsed_result
                print(f"✅ File processed: {len(resume_text)} characters")

        # Ensure we have valid text input
        if not resume_text or not isinstance(resume_text, str):
            resume_text = ""
        if not job_description or not isinstance(job_description, str):
            job_description = ""

        # Strict truncation to prevent timeouts (5K is the sweet spot)
        resume_text = resume_text[:5000]
        job_description = job_description[:5000]
        
        print("🔤 Extracting keywords...")
        # Extract keywords with shorter limits for speed
        job_keywords = bot_audit.extract_keywords(job_description, max_keywords=15) or []
        
        print(f"✅ Found {len(job_keywords)} job keywords")
        
        # Direct matching (faster than nested loops)
        resume_lower = resume_text.lower()
        matched_keywords = [kw for kw in job_keywords if kw in resume_lower] or []
        missing_keywords = [kw for kw in job_keywords if kw not in resume_lower][:8] or []
        
        keyword_score = (len(matched_keywords) / len(job_keywords) * 100) if job_keywords else 0
        
        print("🤖 Analyzing bot safety...")
        # Bot safety analysis
        bot_safety = bot_audit.check_bot_safety(resume_text, structure_audit) or {}
        
        print("📊 Computing advanced metrics...")
        # Advanced metrics
        metrics = bot_audit.analyze_advanced_metrics(resume_text) or {}
        
        # Calculate final ATS score
        ats_score = min(95, max(5, int((keyword_score * 0.6) + (bot_safety.get("safety_score", 0) * 0.4))))
        
        # Determine overall grade
        if ats_score >= 85: 
            overall_grade = "A"
        elif ats_score >= 70: 
            overall_grade = "B"
        elif ats_score >= 55: 
            overall_grade = "C"
        else: 
            overall_grade = "D"
        
        print(f"🎯 Analysis complete: {ats_score}% (Grade {overall_grade})")
        
        # Make sure the RETURN object has every key the frontend expects
        return {
            "success": True,
            "ats_score": ats_score,
            "overall_grade": overall_grade,
            "keyword_analysis": {
                "score": round(keyword_score, 1),
                "matched_keywords": matched_keywords or [], # 'or []' prevents null
                "missing_keywords": missing_keywords or [],
                "total_job_keywords": len(job_keywords)
            },
            "bot_safety": bot_safety or {},
            "metrics": metrics or {},
            "recommendations": {
                "priority_fixes": bot_safety.get("critical_warnings", [])[:3],
                "impact_improvements": [
                    f"Add {max(0, 5-metrics.get('metric_count', 0))} more quantified achievements" if metrics.get('metric_count', 0) < 5 else "Good use of metrics",
                    f"Replace {len(metrics.get('fluff_words', []))} buzzwords with specific examples" if metrics.get('fluff_words', []) else "Content quality is good"
                ],
                "overall_grade": overall_grade
            }
        }
        
    except Exception as e:
        print(f"❌ Job matching error: {e}")
        return {
            "success": False,
            "error": str(e),
            "ats_score": 0,
            "overall_grade": "F",
            "keyword_analysis": {
                "score": 0,
                "matched_keywords": [], # Always array, never null
                "missing_keywords": [],
                "total_job_keywords": 0
            },
            "bot_safety": {
                "safety_score": 0,
                "is_bot_readable": False,
                "critical_warnings": [f"Analysis failed: {str(e)}"],
                "layout_risk": "CRITICAL",
                "found_sections": []
            },
            "metrics": {
                "readability": {"grade_level": 0, "assessment": "Analysis failed"},
                "verb_strength": 0,
                "metric_count": 0,
                "stuffed_keywords": [],
                "fluff_words": [],
                "impact_score": 0,
                "digital_presence": {"linkedin_found": False, "email_found": False},
                "analysis_possible": False
            },
            "recommendations": {
                "priority_fixes": [f"Analysis failed: {str(e)}"],
                "impact_improvements": [],
                "overall_grade": "F"
            }
        }

# ...existing code...
@app.post("/api/generate-resume-file")
async def generate_resume_file(
    resume_text: str = Form(...), 
    format_type: str = Form("pdf")
):
    try:
        temp_dir = tempfile.mkdtemp()
        timestamp = "2026-01-26"
        file_name = f"ATS_Optimized_Resume_{timestamp}.{format_type.lower()}"
        file_path = os.path.join(temp_dir, file_name)

        if format_type.lower() == "pdf":
            doc = SimpleDocTemplate(file_path, pagesize=letter, 
                                  topMargin=72, bottomMargin=72,
                                  leftMargin=72, rightMargin=72)
            
            styles = getSampleStyleSheet()
            story = []
            
            for line in resume_text.split('\n'):
                if line.strip():
                    # Determine appropriate style
                    if line.isupper() or (len(line) < 50 and not line.startswith('•')):
                        story.append(Paragraph(line, styles['Heading2']))
                    else:
                        story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            
        else:  # DOCX
            doc = Document()
            for line in resume_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(file_path)

        return FileResponse(
            file_path, 
            filename=file_name,
            headers={"Content-Disposition": f"attachment; filename={file_name}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File generation error: {str(e)}")

@app.get("/api/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": "2026-01-26",
        "features": {
            "spacy_nlp": nlp is not None,
            "pdf_processing": "extract_pdf_text" in globals() and callable(extract_pdf_text),
            "docx_processing": True,
            "advanced_metrics": True,
            "file_generation": True
        },
        "version": "2.0.0"
    }
# ...existing code...

# --- STARTUP SECTION (KEEP FOR LOCAL DEVELOPMENT) ---
if __name__ == "__main__":
    print("🚀 Starting Trendsetter AI Resume Helper 2026...")
    print(f"📊 spaCy NLP: {'✅ Active' if nlp else '⚠️ Fallback mode'}")
    print(f"📄 PDF Processing: {'✅ Active' if 'extract_pdf_text' in globals() else '⚠️ DOCX only'}")
    print("🎯 Enterprise Analytics: ✅ Active")
    
    uvicorn.run("app:app", host="0.0.0.0", port=8000, log_level="info", reload=True)